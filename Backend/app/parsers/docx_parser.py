"""
DOCX text extractor using python-docx.
Extracts both paragraph text and table cell content.
"""
import logging

from app.core.exceptions import ParseFailureException
from app.parsers.base import BaseParser

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    def extract_text(self, file_path: str) -> str:
        try:
            from docx import Document

            doc = Document(file_path)
            parts: list[str] = []

            # Paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Tables (common in resume templates)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)

            result = "\n".join(parts)
            if not result.strip():
                raise ParseFailureException(
                    message="The DOCX file appears to be empty or contains no readable text.",
                    details={"file_path": file_path},
                )
            return result.strip()

        except ParseFailureException:
            raise
        except Exception as exc:
            logger.exception("DOCX parse failed: %s", exc)
            raise ParseFailureException(
                message=f"Failed to read the DOCX file: {exc}",
                details={"file_path": file_path},
            )
