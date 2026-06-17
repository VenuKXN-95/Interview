"""
Unit tests for resume parsers.
"""
import os
import tempfile

import pytest

from app.parsers.pdf_parser import PDFParser
from app.parsers.docx_parser import DOCXParser
from app.parsers.txt_parser import TXTParser
from app.core.exceptions import ParseFailureException


class TestTXTParser:
    def test_extract_utf8(self, tmp_path):
        p = tmp_path / "resume.txt"
        p.write_text("John Doe\njohn@example.com\nPython, FastAPI", encoding="utf-8")
        parser = TXTParser()
        text = parser.extract_text(str(p))
        assert "John Doe" in text
        assert "john@example.com" in text

    def test_extract_latin1(self, tmp_path):
        p = tmp_path / "resume.txt"
        p.write_bytes("Ren\xe9 Dupont\nrene@email.com".encode("latin-1"))
        parser = TXTParser()
        text = parser.extract_text(str(p))
        assert "Dupont" in text

    def test_empty_file_raises(self, tmp_path):
        p = tmp_path / "empty.txt"
        p.write_text("   \n   ", encoding="utf-8")
        parser = TXTParser()
        with pytest.raises(ParseFailureException):
            parser.extract_text(str(p))


class TestDOCXParser:
    def test_extract_text(self, tmp_path):
        """Create a minimal DOCX and parse it."""
        pytest.importorskip("docx")
        from docx import Document
        doc = Document()
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Skills: Python, Docker")
        docx_path = str(tmp_path / "test.docx")
        doc.save(docx_path)

        parser = DOCXParser()
        text = parser.extract_text(docx_path)
        assert "John Doe" in text
        assert "Python" in text

    def test_table_extraction(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Company"
        table.cell(0, 1).text = "Role"
        table.cell(1, 0).text = "TechCorp"
        table.cell(1, 1).text = "Engineer"
        docx_path = str(tmp_path / "table.docx")
        doc.save(docx_path)

        parser = DOCXParser()
        text = parser.extract_text(docx_path)
        assert "TechCorp" in text
        assert "Engineer" in text
